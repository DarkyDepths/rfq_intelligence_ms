from __future__ import annotations

from pathlib import Path

from docx import Document

from src.services.package_parser.extractors.rvl_extractor import RvlExtractor
from src.services.package_parser.extractors.section_classifier import SectionClassifier
from src.services.package_parser.scanners.tree_scanner import TreeScanner


_PACKAGE_ROOT_NAME = "SA-AYPP-6-MR-022_COLLECTION VESSEL - CDS-REV-00"


def _fixture_package_root() -> Path:
    repo_root = Path(__file__).resolve().parents[2]
    fixture_root = (
        repo_root
        / "local_fixtures"
        / "rfq_created"
        / "source_package_sample_001"
        / _PACKAGE_ROOT_NAME
    )
    if not fixture_root.exists():
        raise FileNotFoundError(f"Fixture package root not found: {fixture_root}")
    return fixture_root


def test_rvl_extractor_parses_real_fixture_docx() -> None:
    fixture_root = _fixture_package_root()
    inventory = TreeScanner().scan(fixture_root)
    registry = SectionClassifier().classify(inventory)

    profile = RvlExtractor().extract(inventory, registry, fixture_root)

    assert profile is not None
    assert profile.source_file == "03-RVL/SA-AYPP-MR-023_ RVL-.docx"
    assert profile.source_format == "docx"
    assert profile.mr_number_in_rvl == "MR-023"
    assert profile.total_vendors >= 10
    assert profile.total_vendors == len(profile.vendors)
    assert "GEYAD FACTORY FOR PRESSURE VESSELS" in profile.unique_vendor_names
    assert "SA" in profile.unique_countries
    assert "DE" in profile.unique_countries
    assert "JP" in profile.unique_countries
    assert "6000002521" in profile.nine_com_codes
    assert profile.unique_vendor_names
    assert profile.unique_countries
    assert profile.nine_com_codes
    assert profile.source_file.endswith(".docx")


def test_rvl_extractor_parses_synthetic_docx_without_sample_specific_assumptions(tmp_path: Path) -> None:
    package_root = tmp_path / "SYNTH-2-MR-009_VENDOR LIST-REV-00"
    rvl_section = package_root / "03-RVL"
    rvl_section.mkdir(parents=True)

    document_path = rvl_section / "vendor_matrix.docx"
    document = Document()
    document.add_paragraph("Synthetic RVL document")
    noise_table = document.add_table(rows=2, cols=2)
    noise_table.rows[0].cells[0].text = "Ignore"
    noise_table.rows[0].cells[1].text = "Me"
    noise_table.rows[1].cells[0].text = "Still"
    noise_table.rows[1].cells[1].text = "Noise"

    table = document.add_table(rows=3, cols=6)
    table.rows[0].cells[0].text = "Country Name"
    table.rows[0].cells[1].text = " Name "
    table.rows[0].cells[2].text = "9com"
    table.rows[0].cells[3].text = "MATERIAL   DESCRIPTION"
    table.rows[0].cells[4].text = "Country Code"
    table.rows[0].cells[5].text = "mfr id"
    table.rows[1].cells[0].text = "United Arab Emirates"
    table.rows[1].cells[1].text = "SYNTH VENDOR"
    table.rows[1].cells[2].text = "7000000007"
    table.rows[1].cells[3].text = "SYNTHETIC EQUIPMENT"
    table.rows[1].cells[4].text = "AE"
    table.rows[1].cells[5].text = "MFR-77"
    table.rows[2].cells[0].text = ""
    table.rows[2].cells[1].text = ""
    table.rows[2].cells[2].text = ""
    table.rows[2].cells[3].text = ""
    table.rows[2].cells[4].text = ""
    table.rows[2].cells[5].text = ""
    document.save(document_path)

    (rvl_section / "vendor_matrix.pdf").write_text("ignore", encoding="utf-8")

    inventory = TreeScanner().scan(package_root)
    registry = SectionClassifier().classify(inventory)

    profile = RvlExtractor().extract(inventory, registry, package_root)

    assert profile is not None
    assert profile.source_file == "03-RVL/vendor_matrix.docx"
    assert profile.source_format == "docx"
    assert profile.total_vendors == 1
    assert profile.unique_vendor_names == ["SYNTH VENDOR"]
    assert profile.unique_countries == ["AE"]
    assert profile.nine_com_codes == ["7000000007"]
    assert profile.mr_number_in_rvl is None

    vendor = profile.vendors[0]
    assert vendor.vendor_name == "SYNTH VENDOR"
    assert vendor.nine_com == "7000000007"
    assert vendor.material_description == "SYNTHETIC EQUIPMENT"
    assert vendor.manufacturer_id == "MFR-77"
    assert vendor.country_code == "AE"
    assert vendor.country_name == "United Arab Emirates"
